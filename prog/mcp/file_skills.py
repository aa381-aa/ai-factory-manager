"""
文件技能模块

文件用途：
    提供AI工厂管家所需的文件读写技能，覆盖PDF、Word、Excel、图片、
    通用文件写入（MinIO/TOS对象存储）五大场景。
    每个技能封装为独立的FileSkill子类，统一通过execute(params)调用。

对应技术规格章节：
    - §1.3 MCP工具中心 - 文件技能（PDF/Word/Excel/图片读写）
    - §1.3.3 MCP技能注册机制（本模块定义可被SkillRegistry注册的技能类）

替代demo：
    替代 demo server.py 中无文件处理能力的缺陷。
    demo的server.py仅依赖data_manager.py直接读写本地JSON，
    无法解析客户上传的PDF报价单、Excel BOM、Word合同、图纸图片等，
    本模块补齐该能力。

依赖：
    - core/file_storage.py：统一文件存储抽象（MinIO/TOS/本地），
      FileWriter子类调用其upload接口完成对象存储写入。
"""

import csv
import io
import os
import tempfile
import traceback
from typing import Any, Dict, List, Optional


class SkillResult:
    """技能执行结果。

    所有FileSkill子类execute均返回该对象，便于上层统一处理成功/失败、
    携带结构化数据与原始文本。
    """

    def __init__(self, success: bool, data: Any = None, text: str = "",
                 error: Optional[str] = None, meta: Optional[Dict[str, Any]] = None) -> None:
        self.success = success
        self.data = data
        self.text = text
        self.error = error
        self.meta = meta or {}

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典（用于JSON序列化返回前端）。"""
        return {
            "success": self.success,
            "data": self.data,
            "text": self.text,
            "error": self.error,
            "meta": self.meta,
        }


class FileSkill:
    """文件技能基类。

    设计意图：
        统一所有文件技能的调用契约（execute(params) -> SkillResult），
        屏蔽底层解析库差异（pdfplumber/python-docx/openpyxl/pytesseract等）。
        子类需实现_execute方法，基类负责异常包裹与结果封装。

    参数：
        name: 技能名称（用于SkillRegistry注册）
        supported_formats: 支持的文件扩展名列表，如 ['.pdf']

    返回值：
        execute返回SkillResult实例。
    """

    name: str = "base"
    supported_formats: list = []

    def __init__(self, name: str = "", supported_formats: Optional[list] = None) -> None:
        self.name = name or self.name
        self.supported_formats = supported_formats if supported_formats is not None else self.supported_formats

    def execute(self, params: Dict[str, Any]) -> SkillResult:
        """执行技能。

        参数：
            params: 技能参数，至少包含 file_path 或 file_url；
                    部分技能需要额外参数（如OCR语言、Excel sheet名）。

        返回：
            SkillResult
        """
        try:
            return self._execute(params)
        except Exception as e:
            # v6.96 P1-12：错误携带完整堆栈，便于定位真实失败点
            return SkillResult(success=False, error=f"{e}\n{traceback.format_exc()}")

    def _execute(self, params: Dict[str, Any]) -> SkillResult:
        """子类实现的具体执行逻辑。"""
        return SkillResult(success=False, error="基类未实现具体逻辑")

    def supports(self, file_path: str) -> bool:
        """判断当前技能是否支持该文件格式。"""
        if "*" in self.supported_formats:
            return True
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.supported_formats


class PDFReader(FileSkill):
    """PDF文件读取技能。

    支持格式：.pdf
    使用场景：
        - 解析客户上传的PDF报价单、采购合同
        - 提取技术规格书的文本与表格
    实现说明：底层使用pdfplumber提取文本与表格，pypdf作为兜底。
    """

    name = "pdf_reader"
    supported_formats = [".pdf"]

    def __init__(self) -> None:
        super().__init__(name=self.name, supported_formats=self.supported_formats)

    def _execute(self, params: Dict[str, Any]) -> SkillResult:
        file_path = params.get("file_path") or params.get("file_url")
        if not file_path:
            return SkillResult(success=False, error="缺少 file_path 参数")

        text_parts: List[str] = []
        tables: List[List[List[str]]] = []

        # 优先使用 pdfplumber 提取文本与表格
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
            return SkillResult(
                success=True,
                text="\n".join(text_parts),
                data={"tables": tables, "pages": len(pdf.pages)},
                meta={"engine": "pdfplumber"},
            )
        except ImportError:
            pass  # pdfplumber 未安装，尝试 pypdf 兜底

        # 使用 pypdf 作为兜底
        try:
            from pypdf import PdfReader as _PdfReader

            reader = _PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return SkillResult(
                success=True,
                text="\n".join(text_parts),
                data={"pages": len(reader.pages)},
                meta={"engine": "pypdf"},
            )
        except ImportError:
            return SkillResult(
                success=False,
                error="PDF解析依赖未安装（pdfplumber 或 pypdf），请先 pip install pdfplumber",
            )


class WordReader(FileSkill):
    """Word文档读取技能。

    支持格式：.docx, .doc
    使用场景：
        - 读取销售合同模板
        - 提取工艺说明文档
    实现说明：底层使用python-docx读取段落与表格。
    """

    name = "word_reader"
    supported_formats = [".docx", ".doc"]

    def __init__(self) -> None:
        super().__init__(name=self.name, supported_formats=self.supported_formats)

    def _execute(self, params: Dict[str, Any]) -> SkillResult:
        file_path = params.get("file_path") or params.get("file_url")
        if not file_path:
            return SkillResult(success=False, error="缺少 file_path 参数")

        try:
            from docx import Document
        except ImportError:
            return SkillResult(
                success=False,
                error="Word解析依赖未安装（python-docx），请先 pip install python-docx",
            )

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables: List[List[List[str]]] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)

        return SkillResult(
            success=True,
            text="\n".join(paragraphs),
            data={"paragraphs": paragraphs, "tables": tables},
            meta={"engine": "python-docx"},
        )


class ExcelReader(FileSkill):
    """Excel表格读取技能。

    支持格式：.xlsx, .xls, .csv
    使用场景：
        - 读取BOM清单
        - 解析库存盘点表、客户对账单
    实现说明：底层使用openpyxl读取xlsx，pandas读取xls/csv，
              自动识别表头与数据区域。
    """

    name = "excel_reader"
    supported_formats = [".xlsx", ".xls", ".csv"]

    def __init__(self) -> None:
        super().__init__(name=self.name, supported_formats=self.supported_formats)

    def _execute(self, params: Dict[str, Any]) -> SkillResult:
        file_path = params.get("file_path") or params.get("file_url")
        if not file_path:
            return SkillResult(success=False, error="缺少 file_path 参数")

        ext = os.path.splitext(file_path)[1].lower()

        # CSV 文件直接用标准库解析
        if ext == ".csv":
            rows = self._parse_csv(file_path)
            return SkillResult(
                success=True,
                data=rows,
                text=f"CSV解析完成，共 {len(rows)} 行",
                meta={"engine": "csv", "format": "csv"},
            )

        # xlsx 文件优先用 openpyxl
        if ext == ".xlsx":
            try:
                import openpyxl

                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                sheet_name = params.get("sheet")
                ws = wb[sheet_name] if sheet_name else wb.active
                rows = self._parse_openpyxl_sheet(ws)
                wb.close()
                return SkillResult(
                    success=True,
                    data=rows,
                    text=f"Excel解析完成，共 {len(rows)} 行",
                    meta={"engine": "openpyxl", "sheet": ws.title},
                )
            except ImportError:
                pass  # openpyxl 未安装，尝试 pandas 兜底

        # 尝试用 pandas 读取（支持 xls / xlsx）
        try:
            import pandas as pd

            df = pd.read_excel(file_path, sheet_name=params.get("sheet", 0))
            records = df.fillna("").to_dict(orient="records")
            # 将非标准类型转为字符串
            for record in records:
                for key, value in record.items():
                    if not isinstance(value, (str, int, float, bool, type(None))):
                        record[key] = str(value)
            return SkillResult(
                success=True,
                data=records,
                text=f"Excel解析完成，共 {len(records)} 行",
                meta={"engine": "pandas"},
            )
        except ImportError:
            return SkillResult(
                success=False,
                error="Excel解析依赖未安装（openpyxl 或 pandas），请先 pip install openpyxl",
            )

    @staticmethod
    def _parse_csv(file_path: str) -> List[Dict[str, Any]]:
        """使用标准库csv模块解析CSV文件。"""
        with open(file_path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            return [dict(row) for row in reader]

    @staticmethod
    def _parse_openpyxl_sheet(ws) -> List[Dict[str, Any]]:
        """解析 openpyxl 工作表，首行作为表头。"""
        rows_iter = ws.iter_rows(values_only=True)
        try:
            headers = [str(h).strip() if h is not None else f"col_{i}"
                       for i, h in enumerate(next(rows_iter))]
        except StopIteration:
            return []
        records: List[Dict[str, Any]] = []
        for row in rows_iter:
            if all(v is None for v in row):
                continue
            record = {}
            for i, value in enumerate(row):
                key = headers[i] if i < len(headers) else f"col_{i}"
                record[key] = value if value is not None else ""
            records.append(record)
        return records


class ImageReader(FileSkill):
    """图片识别技能（OCR）。

    支持格式：.png, .jpg, .jpeg, .bmp, .tiff
    使用场景：
        - 识别图纸版本号、技术参数标注
        - 提取手写质检记录、送货单照片
    实现说明：底层使用pytesseract（多语言OCR）+ Pillow图像预处理，
              可选接入豆包视觉大模型提升复杂场景识别精度。
    """

    name = "image_reader"
    supported_formats = [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]

    def __init__(self) -> None:
        super().__init__(name=self.name, supported_formats=self.supported_formats)

    def _execute(self, params: Dict[str, Any]) -> SkillResult:
        file_path = params.get("file_path") or params.get("file_url")
        if not file_path:
            return SkillResult(success=False, error="缺少 file_path 参数")

        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return SkillResult(
                success=False,
                error="OCR依赖未安装（pytesseract + Pillow），请先 pip install pytesseract Pillow",
            )

        lang = params.get("lang", "chi_sim+eng")
        # v6.96 P1-17：with 上下文管理 PIL 图片句柄，避免文件句柄泄漏
        with Image.open(file_path) as image:
            text = pytesseract.image_to_string(image, lang=lang)
            image_size = image.size
        return SkillResult(
            success=True,
            text=text.strip(),
            data={"lang": lang, "size": image_size},
            meta={"engine": "pytesseract"},
        )


class FileWriter(FileSkill):
    """文件写入技能（对象存储）。

    支持格式：任意（按二进制流写入）
    使用场景：
        - 将生成报表、合同、BOM导出文件写入对象存储
        - 上传用户附件、图纸文件
    实现说明：底层调用 core/file_storage.py 的 S3Storage.upload，
              统一对接 MinIO 或火山引擎TOS，返回可访问URL。
    """

    name = "file_writer"
    supported_formats = ["*"]

    def __init__(self) -> None:
        super().__init__(name=self.name, supported_formats=self.supported_formats)

    def _execute(self, params: Dict[str, Any]) -> SkillResult:
        object_name = params.get("object_name") or params.get("file_path")
        content = params.get("content") or params.get("data")
        if not object_name or content is None:
            return SkillResult(success=False, error="缺少 object_name 或 content 参数")

        # 将内容转为字节流
        if isinstance(content, str):
            content_bytes = content.encode("utf-8")
        elif isinstance(content, bytes):
            content_bytes = content
        else:
            content_bytes = str(content).encode("utf-8")

        try:
            from prog.core.file_storage import get_file_storage

            storage = get_file_storage()
            result_path = storage.upload(
                object_name=object_name,
                data=io.BytesIO(content_bytes),
                length=len(content_bytes),
                content_type=params.get("content_type"),
                metadata=params.get("metadata"),
            )
            url = storage.get_presigned_url(result_path)
            return SkillResult(
                success=True,
                data={"path": result_path, "url": url},
                text=f"文件已上传至 {result_path}",
                meta={"engine": "S3Storage"},
            )
        except Exception as e:
            return SkillResult(success=False, error=f"文件写入失败: {e}")


class FileSkills:
    """文件技能集合类。

    提供文件读写、图纸管理、Excel解析、报告生成等高级文件操作能力。
    当注入 file_storage（S3Storage）时使用对象存储，否则降级为本地文件系统。

    属性：
        _storage: 文件存储实例（S3Storage 或 None）
        _local_base: 本地降级模式的根目录
    """

    def __init__(self, file_storage=None) -> None:
        """初始化文件技能。

        参数：
            file_storage: 文件存储实例（S3Storage）；为 None 时降级为本地文件系统
        """
        self._storage = file_storage
        # 本地降级模式的根目录
        self._local_base = os.path.join(tempfile.gettempdir(), "ai_factory_files")
        os.makedirs(self._local_base, exist_ok=True)

    # ============================================================
    # 通用文件读写
    # ============================================================
    def read_file(self, path: str) -> str:
        """读取文件内容。

        参数：
            path: 文件路径（对象存储路径或本地路径）

        返回：
            文件文本内容字符串
        """
        if self._storage is not None:
            data = self._storage.download(path)
            if isinstance(data, bytes):
                return data.decode("utf-8")
            return str(data)
        # 本地降级
        local_path = self._local_path(path)
        with open(local_path, "r", encoding="utf-8") as f:
            return f.read()

    def write_file(self, path: str, content: str) -> str:
        """写入文件。

        参数：
            path: 文件路径
            content: 文件内容字符串

        返回：
            写入后的文件路径
        """
        if self._storage is not None:
            content_bytes = content.encode("utf-8") if isinstance(content, str) else content
            return self._storage.upload(
                object_name=path,
                data=io.BytesIO(content_bytes),
                length=len(content_bytes),
            )
        # 本地降级
        local_path = self._local_path(path)
        parent = os.path.dirname(local_path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(local_path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    def list_files(self, prefix: str = "") -> list:
        """列举文件。

        参数：
            prefix: 路径前缀过滤

        返回：
            文件信息列表，每项含 name / size 字段
        """
        if self._storage is not None:
            return self._storage.list_objects(prefix=prefix or None)
        # 本地降级
        results = []
        search_root = self._local_path(prefix) if prefix else self._local_base
        if os.path.isfile(search_root):
            results.append({
                "name": prefix,
                "size": os.path.getsize(search_root),
            })
        else:
            for root, _dirs, files in os.walk(search_root):
                for f in files:
                    full_path = os.path.join(root, f)
                    rel_path = os.path.relpath(full_path, self._local_base).replace("\\", "/")
                    results.append({
                        "name": rel_path,
                        "size": os.path.getsize(full_path),
                    })
        return results

    # ============================================================
    # 图纸管理
    # ============================================================
    def upload_drawing(self, product_code: str, file_data: bytes, version: str = "1.0") -> dict:
        """上传图纸。

        参数：
            product_code: 产品编码
            file_data: 图纸文件二进制数据
            version: 图纸版本号

        返回：
            {"success": True, "path": "...", "url": "..."} 或 {"success": False, "error": "..."}
        """
        object_name = f"drawings/{product_code}/{version}"
        try:
            if self._storage is not None:
                path = self._storage.upload(
                    object_name=object_name,
                    data=io.BytesIO(file_data),
                    length=len(file_data),
                )
                url = self._storage.get_presigned_url(path)
                return {"success": True, "path": path, "url": url}
            # 本地降级
            local_path = self._local_path(object_name)
            parent = os.path.dirname(local_path)
            if parent:
                os.makedirs(parent, exist_ok=True)
            with open(local_path, "wb") as f:
                f.write(file_data)
            return {"success": True, "path": object_name, "url": f"file://{os.path.abspath(local_path)}"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def download_drawing(self, product_code: str, version: str = "1.0") -> bytes:
        """下载图纸。

        参数：
            product_code: 产品编码
            version: 图纸版本号

        返回：
            图纸文件二进制数据
        """
        object_name = f"drawings/{product_code}/{version}"
        if self._storage is not None:
            return self._storage.download(object_name)
        # 本地降级
        local_path = self._local_path(object_name)
        with open(local_path, "rb") as f:
            return f.read()

    # ============================================================
    # Excel 解析（csv 降级）
    # ============================================================
    def parse_excel(self, file_data) -> List[Dict[str, Any]]:
        """解析Excel（csv降级）。

        参数：
            file_data: 文件路径字符串，或二进制数据

        返回：
            行记录字典列表，首行作为表头
        """
        # 如果传入的是文件路径，委托给 ExcelReader
        if isinstance(file_data, str) and os.path.isfile(file_data):
            reader = ExcelReader()
            result = reader.execute({"file_path": file_data})
            if result.success:
                return result.data or []
            return []

        # 如果传入的是二进制数据
        if isinstance(file_data, bytes):
            # 先尝试用 openpyxl 从内存读取
            try:
                import openpyxl

                wb = openpyxl.load_workbook(io.BytesIO(file_data), read_only=True, data_only=True)
                ws = wb.active
                rows = ExcelReader._parse_openpyxl_sheet(ws)
                wb.close()
                return rows
            except ImportError:
                pass
            # 降级为 CSV 解析
            try:
                text = file_data.decode("utf-8-sig")
                reader = csv.DictReader(io.StringIO(text))
                return [dict(row) for row in reader]
            except Exception:
                return []

        return []

    # ============================================================
    # 报告生成
    # ============================================================
    def generate_report(self, template: str, data: Dict[str, Any]) -> str:
        """生成报告。

        使用简单模板替换：将 {key} 占位符替换为 data 中对应的值。

        参数：
            template: 模板字符串，含 {key} 占位符
            data: 模板变量字典

        返回：
            替换后的报告字符串
        """
        try:
            return template.format(**data)
        except (KeyError, IndexError):
            # 如果 format 失败（占位符与数据不匹配），使用逐个替换
            result = template
            for key, value in data.items():
                result = result.replace("{" + key + "}", str(value))
            return result

    # ============================================================
    # 内部辅助方法
    # ============================================================
    def _local_path(self, object_name: str) -> str:
        """计算本地降级模式下对象的完整文件路径。

        参数：
            object_name: 对象存储路径

        返回：
            本地文件系统完整路径
        """
        safe_name = object_name.replace("\\", "/")
        path = os.path.join(self._local_base, safe_name)
        # v6.96 P0-3：realpath 归一化 + os.sep 边界，杜绝 ../ 同前缀穿越逃逸
        # （原实现 abspath + startswith 无路径分隔符边界，`../ai_factory_files_evil/x`
        #  前缀同构即可绕过；对齐 server.py _resolve_allowed_file_path 实现）
        base = os.path.realpath(self._local_base)
        real = os.path.realpath(path)
        if not (real == base or real.startswith(base + os.sep)):
            raise ValueError(f"非法的对象路径: {object_name}")
        return real


# ============================================================
# DEBUG 自检（发行版自动跳过，仅在 PROG_DEBUG=1 时执行）
# 验证基座：模块导入、核心类定义、基本结构完整性
# ============================================================
def _self_test():
    """DEBUG模式自检：验证模块基座正确"""
    from prog.core.debug import hello_world

    assert SkillResult is not None, "SkillResult 类未定义"
    assert FileSkill is not None, "FileSkill 类未定义"
    assert FileSkills is not None, "FileSkills 类未定义"
    # 验证 FileSkills 本地降级模式基本功能
    skills = FileSkills(file_storage=None)
    skills.write_file("test_dir/hello.txt", "测试内容")
    content = skills.read_file("test_dir/hello.txt")
    assert content == "测试内容", f"本地文件读写验证失败: {content}"
    files = skills.list_files("test_dir/")
    assert len(files) >= 1, "列举文件验证失败"
    report = skills.generate_report("产品: {product}, 数量: {qty}", {"product": "A-202", "qty": 100})
    assert "A-202" in report and "100" in report, f"报告生成验证失败: {report}"
    hello_world(__name__, "FileSkills 本地降级模式验证通过")


from prog.core.debug import DEBUG

if DEBUG:
    _self_test()
